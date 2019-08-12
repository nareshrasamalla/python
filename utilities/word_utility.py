# -*- coding: utf-8 -*-
"""
This file contains utility functions to Read MS Word document.

"""
import win32com.client

import os
import win32com.client#ability to create and use COM(Microsoft Component Object Model) objects

import docx
import re

def read_document_as_text(file):
    word = win32com.client.Dispatch("Word.Application")
    word.visible = False
    print(file)
    word.Documents.Open(file)
    doc = word.ActiveDocument
    docVar = doc.Range().Text
    print(docVar)
    return docVar


wdFormatPDF = 17  # For saving in pdf format

def convert_to_pdf(input_file, output_file):
    in_file = os.path.abspath(input_file)
    out_file = os.path.abspath(output_file)

    word = win32com.client.Dispatch('Word.Application')  # to parse Microsoft Word documents

    doc = word.Documents.Open(in_file)
    doc.SaveAs(out_file, FileFormat=wdFormatPDF)  # saving input_file as "output_file.pdf"
    doc.Close()
    word.Quit()


# This method is used to replace .dat file content
# @ dat_file : .dat file path
# @ val_map  : dictionary for data replacement
def replace_dat_doc_data(dat_file, **val_map):
    with open(dat_file, 'r') as file:
        contents = file.read()

    for key, value in val_map.items():
        contents = contents.replace(key, value)

    with open(dat_file, 'w') as file:
        file.write(contents)


# This method is used to replace word file content
# @ dat_file : word file path
# @ val_map  : dictionary for data replacement
def replace_word_doc_data(file, **val_map):
    print(file)
    doc = docx.Document(file)

    for key, value in val_map.items():
        searchre = re.compile(key)

        # Getting text from paragraph
        for para in doc.paragraphs:
            inline = para.runs

            for i in range(len(inline)):
                text = inline[i].text

                if text:
                    if searchre.search(text):
                        inline[i].text = re.sub(key, value, text)

        # Getting text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        inline = para.runs

                        for i in range(len(inline)):
                            text = inline[i].text

                            if text:
                                if searchre.search(text):
                                    inline[i].text = re.sub(key, value, text)

    doc.save(file)